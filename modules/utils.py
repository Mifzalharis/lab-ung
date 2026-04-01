from docx import Document
from docx2pdf import convert
import io
import os
import pythoncom
import database as db
import uuid

def generate_pdf(booking_data):
    # Initialize COM for Streamlit thread
    pythoncom.CoInitialize()
    
    # Load template
    doc = Document("FORMAT_PEMINJAMAN LAB.docx")
    
    # Get user details
    user = db.get_user_details(booking_data['username'])
    
    # Format date string
    date_str = str(booking_data['start_date'])
    if booking_data['start_date'] != booking_data['end_date']:
        date_str += f" s/d {booking_data['end_date']}"

    # Map fields to values
    mapping = {
        "Nama Pemohon": user['full_name'],
        "Program Studi": booking_data['program_studi'],
        "Kelas": booking_data['kelas'],
        "Laboratorium": booking_data['lab_name'],
        "Mata Kuliah": booking_data['mata_kuliah'],
        "Dosen Pengampu": booking_data['dosen_pengampu'],
        "Tanggal": date_str,
        "Waktu": f"{booking_data['start_time']} - {booking_data['end_time']}"
    }

    # Iterate paragraphs and replace
    for para in doc.paragraphs:
        for key, value in mapping.items():
            if key in para.text:
                if para.text.strip().startswith(key):
                     para.add_run(f" {value}")
                     break

    # Iterate tables (just in case fields are in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in mapping.items():
                        if key in para.text:
                             if para.text.strip().startswith(key):
                                 para.add_run(f" {value}")
                                 break

    # Save to Temp DOCX - Use UUID to avoid PermissionError if previous file is locked
    unique_id = uuid.uuid4()
    temp_docx = f"temp_booking_{booking_data['id']}_{unique_id}.docx"
    temp_pdf = f"temp_booking_{booking_data['id']}_{unique_id}.pdf"
    
    # Use absolute paths for docx2pdf
    abs_docx = os.path.abspath(temp_docx)
    abs_pdf = os.path.abspath(temp_pdf)
    
    doc.save(abs_docx)
    
    try:
        # Convert to PDF
        convert(abs_docx, abs_pdf)
        
        # Read PDF bytes
        with open(abs_pdf, "rb") as f:
            pdf_bytes = f.read()
            
    except Exception as e:
        # Fallback if conversion fails (e.g. no Word installed)
        print(f"Conversion failed: {e}")
        return None
        
    finally:
        # Clean up
        if os.path.exists(abs_docx):
            os.remove(abs_docx)
        if os.path.exists(abs_pdf):
            os.remove(abs_pdf)
            
    return pdf_bytes
