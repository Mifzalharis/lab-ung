from docx import Document
import sys

try:
    doc = Document("FORMAT_PEMINJAMAN LAB.docx")
    print("--- Paragraphs ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"P: {para.text}")

    print("\n--- Tables ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"Row: {row_text}")

except Exception as e:
    print(f"Error reading docx: {e}")
